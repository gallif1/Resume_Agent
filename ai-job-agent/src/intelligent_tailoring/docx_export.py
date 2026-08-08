"""ATS-friendly DOCX export for tailored resumes.

Linear, parser-safe formatting only — no tables, icons, text boxes, or graphics.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    text = (text or "").strip()
    if not text:
        return
    if level <= 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)


def _add_line(doc: Document, text: str, *, bold: bool = False) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _add_bullet(doc: Document, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def build_tailored_cv_docx(
    tailored_cv: dict[str, Any],
    *,
    name: str = "",
    contact_line: str = "",
    target_role: str = "",
) -> bytes:
    """Render structured tailored CV dict to a clean DOCX byte string."""
    cv = tailored_cv if isinstance(tailored_cv, dict) else {}
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if name:
        _add_heading(doc, name, level=1)
    if contact_line:
        p = doc.add_paragraph(contact_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    title = str(cv.get("professional_title") or "").strip()
    if title:
        _add_line(doc, title, bold=True)
    if target_role:
        _add_line(doc, f"Target Role: {target_role}")

    summary = str(cv.get("summary") or cv.get("professional_summary") or "").strip()
    if summary:
        _add_heading(doc, "Professional Summary", level=2)
        _add_line(doc, summary)

    experience = [e for e in (cv.get("experience") or []) if isinstance(e, dict)]
    if experience:
        _add_heading(doc, "Experience", level=2)
        for entry in experience:
            role = str(entry.get("title") or "").strip()
            company = str(entry.get("company") or "").strip()
            dates = str(entry.get("dates") or "").strip()
            heading = role or company or "Experience"
            _add_line(doc, heading, bold=True)
            meta = " | ".join(x for x in (company if role else "", dates) if x)
            if meta:
                _add_line(doc, meta)
            for bullet in entry.get("bullets") or []:
                _add_bullet(doc, str(bullet))

    projects = [p for p in (cv.get("projects") or []) if isinstance(p, dict)]
    if projects:
        _add_heading(doc, "Projects", level=2)
        for entry in projects:
            pname = str(entry.get("name") or "").strip() or "Project"
            _add_line(doc, pname, bold=True)
            desc = str(entry.get("description") or "").strip()
            if desc:
                _add_line(doc, desc)
            for bullet in entry.get("bullets") or []:
                _add_bullet(doc, str(bullet))

    skills = [str(s).strip() for s in (cv.get("skills") or []) if str(s).strip()]
    if skills:
        _add_heading(doc, "Skills", level=2)
        for row in skills:
            _add_line(doc, row)

    try:
        from intelligent_tailoring.canonical_resume import (
            format_education_entry,
            looks_like_raw_data,
            normalize_education_entries,
        )

        education = normalize_education_entries(cv.get("education"))
    except Exception:
        education = [e for e in (cv.get("education") or []) if isinstance(e, dict)]
        looks_like_raw_data = lambda t: False  # noqa: E731
        format_education_entry = None  # type: ignore

    if education:
        rendered_any = False
        for entry in education:
            if format_education_entry is not None:
                formatted = format_education_entry(entry)
                degree = formatted.get("degree") or ""
                institution = formatted.get("institution") or ""
                dates = formatted.get("dates") or ""
                heading = formatted.get("heading") or degree or institution
            else:
                degree = str(entry.get("degree") or "").strip()
                institution = str(entry.get("institution") or "").strip()
                dates = str(entry.get("dates") or "").strip()
                heading = degree or institution
            if not heading or looks_like_raw_data(heading):
                continue
            if looks_like_raw_data(degree) or looks_like_raw_data(institution):
                continue
            if not rendered_any:
                _add_heading(doc, "Education", level=2)
                rendered_any = True
            _add_line(doc, heading, bold=True)
            meta = " | ".join(x for x in (institution if degree else "", dates) if x)
            if meta and not looks_like_raw_data(meta):
                _add_line(doc, meta)

    certs = cv.get("certifications") or []
    if certs:
        _add_heading(doc, "Certifications", level=2)
        for cert in certs:
            if isinstance(cert, dict):
                label = str(cert.get("name") or cert.get("title") or "").strip()
            else:
                label = str(cert).strip()
            if label:
                _add_bullet(doc, label)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def markdown_to_plain_sections(markdown: str) -> dict[str, Any]:
    """Minimal fallback: wrap markdown body as a single summary if structured CV missing."""
    return {
        "professional_title": "",
        "summary": (markdown or "").strip(),
        "skills": [],
        "experience": [],
        "projects": [],
        "education": [],
        "certifications": [],
    }

"""DOCX rendering for tailored CV structured data."""

from __future__ import annotations

import logging
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from cv_tailor.models import TailoredCvData

logger = logging.getLogger("cv_tailor.renderer")


def _add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    text = (text or "").strip()
    if not text:
        return
    if level <= 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)


def _add_line(doc: Document, text: str, *, bold: bool = False) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _add_bullet(doc: Document, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def render_tailored_cv_docx(cv: TailoredCvData) -> bytes:
    """Render structured tailored CV to DOCX bytes."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if cv.name:
        _add_heading(doc, cv.name, level=1)
    if cv.contact:
        p = doc.add_paragraph(cv.contact)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    if cv.summary:
        _add_heading(doc, "Professional Summary", level=2)
        _add_line(doc, cv.summary)

    if cv.experience:
        _add_heading(doc, "Experience", level=2)
        for entry in cv.experience:
            heading = entry.role or entry.company or "Experience"
            _add_line(doc, heading, bold=True)
            meta = " | ".join(
                part for part in (entry.company if entry.role else "", entry.dates) if part
            )
            if meta:
                _add_line(doc, meta)
            for bullet in entry.bullets:
                _add_bullet(doc, bullet)

    if cv.projects:
        _add_heading(doc, "Projects", level=2)
        for project in cv.projects:
            if project.name:
                _add_line(doc, project.name, bold=True)
            if project.description:
                _add_line(doc, project.description)
            for bullet in project.bullets:
                _add_bullet(doc, bullet)

    if cv.skills:
        _add_heading(doc, "Skills", level=2)
        _add_line(doc, ", ".join(cv.skills))

    if cv.education:
        _add_heading(doc, "Education", level=2)
        for edu in cv.education:
            heading = edu.degree or edu.institution
            if heading:
                _add_line(doc, heading, bold=True)
            meta = " | ".join(
                part for part in (edu.institution if edu.degree else "", edu.dates) if part
            )
            if meta:
                _add_line(doc, meta)

    if cv.certifications:
        _add_heading(doc, "Certifications", level=2)
        for cert in cv.certifications:
            _add_bullet(doc, cert)

    buffer = BytesIO()
    doc.save(buffer)
    logger.info("DOCX generated (%d bytes)", len(buffer.getvalue()))
    return buffer.getvalue()

"""Read resume files in multiple formats (PDF, DOCX, TXT, images)."""

from __future__ import annotations

import io
import warnings
from pathlib import Path
from typing import Callable

from config import CV_PATH, RESUMES_DIR

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp")
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
PREFERRED_NAMES = ("cv", "resume", "קורות חיים")


def find_resume_path(preferred: Path = CV_PATH) -> Path | None:
    """Return the best resume file in resumes/ (prefers cv.pdf, then other formats)."""
    if preferred.exists():
        return preferred

    if not RESUMES_DIR.exists():
        return None

    candidates: list[Path] = []
    for path in RESUMES_DIR.iterdir():
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if path.name.startswith("."):
            continue
        candidates.append(path)

    if not candidates:
        return None

    def sort_key(path: Path) -> tuple[int, int, str]:
        stem = path.stem.lower()
        name_rank = next(
            (i for i, name in enumerate(PREFERRED_NAMES) if name in stem),
            len(PREFERRED_NAMES),
        )
        ext_rank = SUPPORTED_EXTENSIONS.index(path.suffix.lower())
        return (name_rank, ext_rank, path.name.lower())

    return sorted(candidates, key=sort_key)[0]


def diagnose_pdf(cv_path: Path) -> list[str]:
    """Return human-readable issues detected in a PDF file."""
    issues: list[str] = []
    if not cv_path.exists():
        return ["File not found"]

    data = cv_path.read_bytes()
    if not data.startswith(b"%PDF"):
        issues.append("The file is not a valid PDF")

    text, backend = _extract_pdf_text(cv_path)
    if not text.strip():
        issues.append("Could not extract text from the PDF (it may be corrupt, scanned, or saved incorrectly)")

    try:
        import fitz

        doc = fitz.open(cv_path)
        if doc.page_count == 0:
            issues.append("The PDF contains no pages")
        else:
            page = doc[0]
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            non_white = sum(
                1
                for y in range(0, pix.height, 40)
                for x in range(0, pix.width, 40)
                if pix.pixel(x, y) != (255, 255, 255)
            )
            if non_white == 0 and not text.strip():
                issues.append(
                    "The page looks blank - the PDF is probably corrupt (e.g. from OneDrive sync). "
                    "Re-export from Word/Google Docs and save again."
                )
    except Exception:
        pass

    if backend:
        issues.append(f"Last extraction attempt: {backend}")

    return issues


def _extract_pdf_text(cv_path: Path) -> tuple[str, str]:
    """Try multiple PDF backends. Returns (text, backend_name)."""
    backends: list[tuple[str, Callable[[], str]]] = []

    def via_pymupdf() -> str:
        import fitz

        doc = fitz.open(cv_path)
        return "\n".join(page.get_text() for page in doc)

    def via_pdfplumber() -> str:
        import pdfplumber

        with pdfplumber.open(cv_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    def via_pypdf() -> str:
        from pypdf import PdfReader

        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            reader = PdfReader(str(cv_path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

    backends.extend([("pymupdf", via_pymupdf), ("pdfplumber", via_pdfplumber), ("pypdf", via_pypdf)])

    last_error = ""
    for name, fn in backends:
        try:
            text = fn() or ""
            if text.strip():
                return text, name
            last_error = name
        except Exception as exc:
            last_error = f"{name}: {exc}"

    return "", last_error


def _extract_docx_text(cv_path: Path) -> str:
    from docx import Document

    doc = Document(str(cv_path))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_plain_text(cv_path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1255", "latin-1"):
        try:
            return cv_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return cv_path.read_text(encoding="utf-8", errors="replace")


def render_pdf_pages(cv_path: Path, *, max_pages: int = 3, zoom: float = 2.0) -> list[bytes]:
    """Render PDF pages to PNG bytes for OpenAI Vision (scanned resumes)."""
    import fitz

    doc = fitz.open(cv_path)
    images: list[bytes] = []
    matrix = fitz.Matrix(zoom, zoom)

    for page_index in range(min(doc.page_count, max_pages)):
        pix = doc[page_index].get_pixmap(matrix=matrix)
        images.append(pix.tobytes("png"))

    return images


def load_image_bytes(cv_path: Path) -> list[bytes]:
    """Load an image resume file as bytes."""
    return [cv_path.read_bytes()]


def page_images_have_content(image_pages: list[bytes]) -> bool:
    """True when at least one rendered page is not a blank white image."""
    try:
        from PIL import Image
    except ImportError:
        return bool(image_pages)

    for data in image_pages:
        img = Image.open(io.BytesIO(data)).convert("RGB")
        pixels = list(img.getdata())
        if not pixels:
            continue
        non_white = sum(1 for r, g, b in pixels[::500] if (r, g, b) != (255, 255, 255))
        if non_white > 0:
            return True
    return False


def extract_text_from_resume(cv_path: Path) -> tuple[str, str]:
    """Extract text from a resume file. Returns (text, source_label)."""
    if not cv_path.exists():
        return "", "missing"

    ext = cv_path.suffix.lower()

    if ext == ".pdf":
        text, backend = _extract_pdf_text(cv_path)
        return text, f"pdf:{backend}"

    if ext in (".docx", ".doc"):
        if ext == ".doc":
            raise ValueError(
                "Old .doc format is not supported - save from Word as .docx or .pdf"
            )
        return _extract_docx_text(cv_path), "docx"

    if ext in (".txt", ".md"):
        return _extract_plain_text(cv_path), ext.lstrip(".")

    if ext in IMAGE_EXTENSIONS:
        return "", f"image:{ext}"

    return "", "unsupported"
